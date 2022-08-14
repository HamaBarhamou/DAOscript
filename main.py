#!/usr/bin/python3

from docx import Document
from modules.load_file import loadfile

def main():
    document = Document()
    document.add_paragraph('page de gardes')
    document.add_page_break()
    doc = loadfile('text.txt')

    len_table = len(doc)
    table = document.add_table(rows=len_table, cols= 2)

    for i in range(len_table):
        cell1 = table.cell(i, 0)
        cell2 = table.cell(i, 1)
        cell1.text = str(i+1)
        cell2.text = doc[i].upper()

    document.add_page_break()

    for line in doc:
        for _ in range(7):
            document.add_paragraph('\n')
        document.add_paragraph(line.upper())
        document.add_page_break()

    path = "Documents/" + input("Donner un nom au documents .docx: ") + ".docx"
    document.save(path)
    print("Chemin d'acces a votre document: ",path)
main()