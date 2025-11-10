#!/usr/bin/env python3
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import defaultdict


def ajouter_element_sommaire(document, titre_affiche, niveau, numerotation):
    """Ajoute une entrée numérotée au sommaire avec mise en forme élégante."""
    paragraphe = document.add_paragraph()
    paragraphe.paragraph_format.left_indent = Pt(niveau * 15)
    paragraphe.paragraph_format.space_before = Pt(4)
    paragraphe.paragraph_format.space_after = Pt(1)
    run = paragraphe.add_run(f"{numerotation} {titre_affiche}")
    run.font.size = Pt(11)
    run.bold = True
    return f"{numerotation} {titre_affiche}"


def ajouter_intercalaire_centre(document, texte):
    """Ajoute une page de garde centrée horizontalement et verticalement."""
    document.add_page_break()

    # Réduire les marges pour mieux centrer verticalement
    section = document.sections[-1]
    section.top_margin = Inches(3)
    section.bottom_margin = Inches(3)

    # Ajouter un paragraphe centré
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texte)
    run.font.size = Pt(24)
    run.bold = True


def generer_sommaire_dossiers(chemin_dossier, nom_fichier_output="structure_dossier_dossiers_seulement.docx"):
    if not os.path.isdir(chemin_dossier):
        raise ValueError("❌ Le chemin fourni n'est pas un dossier valide.")

    document = Document()
    document.add_heading('📁 Sommaire hiérarchique des dossiers', level=1)

    sections = []
    compteur = defaultdict(int)

    def get_numerotation(niveau):
        return ".".join(str(compteur[i]) for i in range(niveau + 1))

    # Racine
    compteur[0] += 1
    ligne_sommaire = ajouter_element_sommaire(document, ".", 0, get_numerotation(0))
    sections.append(ligne_sommaire)

    for dossier_courant, sous_dossiers, _ in os.walk(chemin_dossier):
        rel_path = os.path.relpath(dossier_courant, chemin_dossier)
        niveau = rel_path.count(os.sep) + 1 if rel_path != '.' else 0
        nom_dossier = os.path.basename(dossier_courant.rstrip(os.sep))

        sous_dossiers.sort(key=str.casefold)

        if rel_path != '.':
            compteur[niveau] += 1
            for deeper in list(compteur.keys()):
                if deeper > niveau:
                    compteur.pop(deeper)
            titre_affiche = nom_dossier + "/"
            ligne_sommaire = ajouter_element_sommaire(document, titre_affiche, niveau, get_numerotation(niveau))
            sections.append(ligne_sommaire)

    # Intercalaires
    for titre_complet in sections:
        ajouter_intercalaire_centre(document, titre_complet)

    document.save(nom_fichier_output)
    print(f"\n✅ Document généré avec succès : {nom_fichier_output}")


# === Programme principal ===
if __name__ == "__main__":
    chemin = input("Entrez le chemin d'accès du dossier à analyser : ").strip()
    try:
        generer_sommaire_dossiers(chemin)
    except Exception as e:
        print(f"❌ Erreur : {e}")

