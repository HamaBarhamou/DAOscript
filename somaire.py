#!/usr/bin/env python3
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import defaultdict


def ajouter_element_sommaire(document, titre_affiche, niveau, numerotation):
    """Ajoute une entrée numérotée au sommaire."""
    paragraphe = document.add_paragraph()
    paragraphe.paragraph_format.left_indent = Pt(niveau * 15)
    run = paragraphe.add_run(f"{numerotation} {titre_affiche}")
    run.font.size = Pt(10.5)
    run.bold = True
    paragraphe.paragraph_format.space_after = Pt(0)
    return f"{numerotation} {titre_affiche}"  # On retourne la ligne exacte pour l'intercalaire


def ajouter_intercalaire(document, texte):
    """Crée une page centrée verticalement et horizontalement avec un titre."""
    document.add_page_break()

    section = document.sections[-1]
    section.top_margin = Inches(3)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texte)
    run.font.size = Pt(20)
    run.bold = True


def lire_fichier_si_texte(chemin_fichier):
    """Lit le fichier texte si possible (ignore les binaires)."""
    try:
        with open(chemin_fichier, 'r', encoding='utf-8') as f:
            return f.read()
    except:
        return None


def generer_document_structure(chemin_dossier, nom_fichier_output="structure_dossier.docx"):
    if not os.path.isdir(chemin_dossier):
        raise ValueError("❌ Le chemin fourni n'est pas un dossier valide.")

    document = Document()
    document.add_heading('📁 Sommaire de la structure du dossier', level=1)

    sections = []  # Liste des titres complets du sommaire (numérotés) + contenu optionnel
    compteur = defaultdict(int)

    def get_numerotation(niveau):
        return ".".join(str(compteur[i]) for i in range(niveau + 1))

    # Racine
    compteur[0] += 1
    ligne_sommaire = ajouter_element_sommaire(document, ".", 0, get_numerotation(0))
    sections.append((ligne_sommaire, None))

    for dossier_courant, sous_dossiers, fichiers in os.walk(chemin_dossier):
        rel_path = os.path.relpath(dossier_courant, chemin_dossier)
        niveau = rel_path.count(os.sep) + 1 if rel_path != '.' else 0
        nom_dossier = os.path.basename(dossier_courant.rstrip(os.sep))

        sous_dossiers.sort(key=str.casefold)
        fichiers.sort(key=str.casefold)

        if rel_path != '.':
            compteur[niveau] += 1
            for deeper in list(compteur.keys()):
                if deeper > niveau:
                    compteur.pop(deeper)
            titre_affiche = nom_dossier + "/"
            ligne_sommaire = ajouter_element_sommaire(document, titre_affiche, niveau, get_numerotation(niveau))
            sections.append((ligne_sommaire, None))

        for fichier in fichiers:
            compteur[niveau + 1] += 1
            for deeper in list(compteur.keys()):
                if deeper > niveau + 1:
                    compteur.pop(deeper)
            chemin_fichier = os.path.join(dossier_courant, fichier)
            contenu = lire_fichier_si_texte(chemin_fichier)
            chemin_affiche = os.path.join(rel_path, fichier) if rel_path != '.' else fichier
            ligne_sommaire = ajouter_element_sommaire(document, fichier, niveau + 1, get_numerotation(niveau + 1))
            sections.append((ligne_sommaire, contenu))

    # Génération des pages : une par section du sommaire
    for titre_complet, contenu in sections:
        ajouter_intercalaire(document, titre_complet)
        if contenu:
            p = document.add_paragraph()
            p.add_run(contenu).font.size = Pt(10.5)

    document.save(nom_fichier_output)
    print(f"\n✅ Le document a été généré avec succès : {nom_fichier_output}")


# === Programme principal ===
if __name__ == "__main__":
    chemin = input("Entrez le chemin d'accès du dossier à analyser : ").strip()
    try:
        generer_document_structure(chemin)
    except Exception as e:
        print(f"❌ Erreur : {e}")

